import streamlit as st
import math
import matplotlib.pyplot as plt
import numpy as np
from datetime import datetime
import pandas as pd
import io
from docx import Document
import base64

# Set page configuration
st.set_page_config(page_title="Elevator Earthpit Calculator", layout="wide")

# Title and Introduction
st.title("🛠️ Elevator Earthpit Calculator")
st.markdown("""
This tool helps calculate the recommended earthpit dimensions and components for elevator grounding systems,
based on international and Indian standards such as IEC 62305-3, IEEE Std 80/81, and IS 3043.
""")

# Sidebar for standards reference
with st.sidebar:
    st.header("Standards Reference")
    st.markdown("""
    - **IEC 62305-3:2011**: Earth resistance ≤ 10 ohms
    - **IEEE Std 80 & 81**: Soil resistivity measurement methods
    - **IS 3043:1987 & 2018**: Indian earthing practices
    """)

# Input section
st.header("Input Parameters")

# First row of inputs
col1, col2, col3 = st.columns(3)

with col1:
    pit_type = st.selectbox("Select Pit Type", [
        "Regular Pit",
        "Chemical Earthing Pit",
        "Maintenance-Free Pit"
    ])
    
with col2:
    plate_material = st.selectbox("Select Plate Material", ["GI (Galvanized Iron)", "Copper"])
    
with col3:
    soil_type = st.selectbox("Soil Type", [
        "Loamy Soil",
        "Clay Soil",
        "Sandy Soil",
        "Rocky Soil",
        "Black Cotton Soil"
    ])

# Second row of inputs
col4, col5, col6 = st.columns(3)

with col4:
    resistivity = st.number_input("Soil Resistivity (Ω·m)", 
                                min_value=1.0, 
                                value=100.0,
                                help="Typical range: Loamy: 50-100, Clay: 100-200, Sandy: 200-1000, Rocky: 1000+")
with col5:
    desired_resistance = st.number_input("Desired Earth Resistance (Ω)", 
                                       min_value=0.1, 
                                       value=5.0,
                                       help="As per standards, should be ≤ 10 ohms")
with col6:
    strip_length = st.number_input("Strip Length (m)", 
                                  min_value=1.0, 
                                  value=10.0,
                                  help="Connecting strip length from pit to equipment")

# Calculation section
st.header("📐 Calculated Outputs")

# Material properties
if plate_material == "GI (Galvanized Iron)":
    plate_size = "500 mm x 500 mm x 10 mm"
    material_factor = 1.0  # Base factor for GI
    strip_type = "GI"
    conductivity = "16%"  # Relative to copper
else:  # Copper
    plate_size = "600 mm x 600 mm x 3 mm"
    material_factor = 0.5  # Copper has better conductivity
    strip_type = "Copper"
    conductivity = "100%"

# Soil type factors
soil_factors = {
    "Loamy Soil": 1.0,
    "Clay Soil": 1.2,
    "Sandy Soil": 1.5,
    "Rocky Soil": 2.0,
    "Black Cotton Soil": 0.8
}

# Pit type configurations
pit_configs = {
    "Regular Pit": {
        "depth_factor": 1.0,
        "filling": "Coal/Salt/Sand mixture in 1:1:1 ratio",
        "lifetime": "3-5 years",
        "maintenance": "Regular watering and periodic maintenance required"
    },
    "Chemical Earthing Pit": {
        "depth_factor": 0.7,  # Chemical treatment improves conductivity
        "filling": "Conductive chemical compound with minerals",
        "lifetime": "10-15 years",
        "maintenance": "Periodic chemical refilling every 5-7 years"
    },
    "Maintenance-Free Pit": {
        "depth_factor": 0.8,
        "filling": "Highly conductive copper oxide compound",
        "lifetime": "15-20 years",
        "maintenance": "No regular maintenance required"
    }
}

# Calculate pit depth considering all factors
base_depth = (resistivity * material_factor * soil_factors[soil_type]) / (4 * math.pi * desired_resistance)
pit_depth = round(base_depth * pit_configs[pit_type]["depth_factor"], 2)

# Display calculations and recommendations
st.subheader("📏 Dimensions and Materials")
st.markdown(f"**Recommended Earthpit Depth:** `{pit_depth} meters`")
st.markdown(f"**Standard {plate_material} Plate Size:** `{plate_size}`")
st.markdown(f"**{strip_type} Strip Required:** `{strip_length} meters`")

st.subheader("⚡ Electrical Properties")
st.markdown(f"**Material Conductivity:** `{conductivity}`")
st.markdown(f"**Soil Type Factor:** `{soil_factors[soil_type]}`")
st.markdown(f"**Expected Resistance:** `{round(desired_resistance * pit_configs[pit_type]['depth_factor'], 2)} Ω`")

st.subheader("🏗️ Pit Specifications")
st.markdown(f"**Pit Type:** `{pit_type}`")
st.markdown(f"**Filling Material:** `{pit_configs[pit_type]['filling']}`")
st.markdown(f"**Expected Lifetime:** `{pit_configs[pit_type]['lifetime']}`")
st.markdown(f"**Maintenance Requirements:** `{pit_configs[pit_type]['maintenance']}`")


# Visual Representation
st.subheader("🎯 Visual Representation")
col_vis1, col_vis2 = st.columns([2, 1])

with col_vis1:
    # Create a visual representation of the earthing pit
    st.markdown(f"""
    ```
    Ground Level
    ─────────────────────────────────────────
    │                                       │
    │   Strip ({strip_type})                │
    │   Length: {strip_length}m             │
    │                                       │
    │   ┌───────────────────┐              │
    │   │    {plate_material}    │          │
    │   │    Plate          │ {pit_depth}m  │
    │   │    {plate_size}   │              │
    │   └───────────────────┘              │
    │                                       │
    │   Filling: {pit_configs[pit_type]['filling']}
    │                                       │
    └───────────────────────────────────────┘
    ```
    """)

with col_vis2:
    # Display a checklist of installed components
    st.markdown("**Installation Checklist:**")
    components = {
        "Earthing Plate": True,
        f"{strip_type} Strip": True,
        "Filling Material": True,
        "Moisture Control": pit_type != "Maintenance-Free Pit",
        "Chemical Compound": pit_type == "Chemical Earthing Pit",
        "Inspection Chamber": True
    }
    
    for component, installed in components.items():
        st.markdown(f"{'✅' if installed else '❌'} {component}")

# Additional Technical Details
st.subheader("🔧 Technical Specifications")
tech_col1, tech_col2 = st.columns(2)

with tech_col1:
    st.markdown("**Material Properties:**")
    material_props = {
        "GI (Galvanized Iron)": {
            "Tensile Strength": "380-550 MPa",
            "Zinc Coating": "610 g/m²",
            "Thickness": "10 mm",
            "Life Expectancy": "15-20 years"
        },
        "Copper": {
            "Tensile Strength": "220-250 MPa",
            "Purity": "99.9%",
            "Thickness": "3 mm",
            "Life Expectancy": "30-50 years"
        }
    }
    for prop, value in material_props[plate_material].items():
        st.markdown(f"- {prop}: `{value}`")

with tech_col2:
    st.markdown("**Pit Environment:**")
    environment_factors = {
        "pH Value": "6.5-7.5",
        "Moisture Content": "10-15%",
        "Temperature Range": "10-45°C",
        "Backfill Resistivity": "< 50 Ω·m"
    }
    for factor, value in environment_factors.items():
        st.markdown(f"- {factor}: `{value}`")

# Weather Impact Analysis
st.subheader("🌦️ Weather Impact Analysis")
weather_impact = {
    "Rainy Season": "Optimal performance, natural moisture maintenance",
    "Summer": "May require additional watering" if pit_type == "Regular Pit" else "Stable performance",
    "Winter": "Slightly increased resistance due to lower ground temperature",
    "Monsoon": "Enhanced performance, monitor for flooding"
}

for season, impact in weather_impact.items():
    st.markdown(f"**{season}:** {impact}")

# Inspection Schedule Generator
st.subheader("📅 Recommended Inspection Schedule")
# Create inspection activities dataframe
basic_activities = [
    {"Activity": "Visual Inspection", "Frequency": "Every inspection", "Notes": "Check for physical damage"},
    {"Activity": "Resistance Measurement", "Frequency": "Every inspection", "Notes": "Should be ≤ 10 ohms"},
    {"Activity": "Connection Check", "Frequency": "Every inspection", "Notes": "Verify all connections are tight"}
]

if pit_type == "Regular Pit":
    additional_activities = [
        {"Activity": "Moisture Check", "Frequency": "Monthly", "Notes": "Maintain adequate moisture"},
        {"Activity": "Salt/Coal Level Check", "Frequency": "Quarterly", "Notes": "Replenish if needed"}
    ]
elif pit_type == "Chemical Earthing Pit":
    additional_activities = [
        {"Activity": "Chemical Compound Check", "Frequency": "Every 2 years", "Notes": "Check compound conductivity"}
    ]
else:
    additional_activities = []

inspection_df = pd.DataFrame(basic_activities + additional_activities)

# Display the inspection table
st.markdown("### Inspection Schedule and Activities")
st.dataframe(inspection_df, use_container_width=True)

# Download buttons
def to_excel():
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        inspection_df.to_excel(writer, sheet_name='Inspection Schedule', index=False)
    return output.getvalue()

def to_word():
    doc = Document()
    doc.add_heading('Earthpit Inspection Schedule', 0)
    
    # Add pit specifications
    doc.add_heading('Pit Specifications', level=1)
    doc.add_paragraph(f'Pit Type: {pit_type}')
    doc.add_paragraph(f'Plate Material: {plate_material}')
    doc.add_paragraph(f'Soil Type: {soil_type}')
    
    # Add inspection table
    doc.add_heading('Inspection Schedule', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Activity'
    header_cells[1].text = 'Frequency'
    header_cells[2].text = 'Notes'
    
    # Add data rows
    for _, row in inspection_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['Activity']
        row_cells[1].text = row['Frequency']
        row_cells[2].text = row['Notes']
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

col1, col2 = st.columns(2)
with col1:
    excel_data = to_excel()
    st.download_button(
        label="📥 Download Excel Report",
        data=excel_data,
        file_name=f"earthpit_inspection_{pit_type.lower().replace(' ', '_')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

with col2:
    word_data = to_word()
    st.download_button(
        label="📄 Download Word Report",
        data=word_data,
        file_name=f"earthpit_inspection_{pit_type.lower().replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Maintenance Guidelines
st.header("🧰 Maintenance Guidelines")

if pit_type == "Regular Pit":
    st.markdown("""
    #### Regular Pit Maintenance:
    - 🌧️ Regular watering (weekly in dry season)
    - 🧂 Annual replenishment of salt/charcoal mixture
    - 🔍 Quarterly inspection for corrosion
    - 📏 Bi-annual earth resistance measurement
    - ⚡ Tightening of connections annually
    """)
elif pit_type == "Chemical Earthing Pit":
    st.markdown("""
    #### Chemical Pit Maintenance:
    - 🧪 Chemical compound check every 2 years
    - 🔄 Refilling of chemical compound every 5-7 years
    - 🔍 Annual inspection of connections
    - 📏 Annual earth resistance measurement
    - 💧 No regular watering required
    """)
else:  # Maintenance-Free Pit
    st.markdown("""
    #### Maintenance-Free Pit Care:
    - 🔍 Annual visual inspection
    - 📏 Annual earth resistance measurement
    - ⚡ Check connections every 2 years
    - 🛡️ No chemical refilling required
    - 💧 No watering required
    """)

# Footer
st.markdown("---")
st.markdown("For assistance write to services@serinthomas.in")